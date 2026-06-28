import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def convert_py_to_docx():
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Document Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Source Code Listing: generate_report.py")
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(26, 76, 128)

    # Reading the generate_report.py code file
    py_file_path = "c:\\Users\\piyush\\OneDrive\\Desktop\\Task Tracker\\generate_report.py"
    if not os.path.exists(py_file_path):
        print(f"Error: {py_file_path} not found.")
        return

    with open(py_file_path, "r", encoding="utf-8") as f:
        code_content = f.read()

    # Split code content by lines
    lines = code_content.splitlines()

    # Create a single-column table to act as a shaded code box
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.9)

    cell = table.cell(0, 0)
    set_cell_background(cell, "F5F5F5")  # Light gray background shading
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

    # Add the lines of code inside the table cell
    first_p = cell.paragraphs[0]
    first_p.paragraph_format.space_before = Pt(0)
    first_p.paragraph_format.space_after = Pt(0)
    first_p.paragraph_format.line_spacing = 1.0

    # Write the first line into the default paragraph in the cell
    if lines:
        run_line = first_p.add_run(lines[0])
        run_line.font.name = 'Consolas'
        run_line.font.size = Pt(8.5)
        run_line.font.color.rgb = RGBColor(40, 40, 40)

    # Write the remaining lines into new paragraphs
    for line in lines[1:]:
        p_line = cell.add_paragraph()
        p_line.paragraph_format.space_before = Pt(0)
        p_line.paragraph_format.space_after = Pt(0)
        p_line.paragraph_format.line_spacing = 1.0
        
        run_line = p_line.add_run(line)
        run_line.font.name = 'Consolas'
        run_line.font.size = Pt(8.5)
        run_line.font.color.rgb = RGBColor(40, 40, 40)

    # Save the output file
    output_path = "c:\\Users\\piyush\\OneDrive\\Desktop\\Task Tracker\\generate_report_code.docx"
    doc.save(output_path)
    print(f"Successfully converted source code to {output_path}")

if __name__ == "__main__":
    convert_py_to_docx()
